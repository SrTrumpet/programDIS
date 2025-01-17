from docx import Document

def docx_to_html(docx_file, html_file):
    # Leer el archivo DOCX
    document = Document(docx_file)

    # Crear un archivo HTML
    with open(html_file, 'w', encoding='utf-8') as html:
        html.write("<html>\n<head>\n<title>Document</title>\n</head>\n<body>\n")

        # Convertir el contenido del DOCX en HTML
        for paragraph in document.paragraphs:
            html.write(f"<p>{paragraph.text}</p>\n")

        # Manejar tablas
        for table in document.tables:
            html.write("<table border='1'>\n")
            for row in table.rows:
                html.write("<tr>\n")
                for cell in row.cells:
                    html.write(f"<td>{cell.text}</td>\n")
                html.write("</tr>\n")
            html.write("</table>\n")

        html.write("</body>\n</html>")

# Uso del script
docx_file_path = "Manual de uso DIS - Gestor de Documentaicon.docx"  # Ruta al archivo .docx
html_file_path = "archivo.html"  # Ruta de salida para el archivo .html

docx_to_html(docx_file_path, html_file_path)
print(f"Archivo HTML generado en: {html_file_path}")
